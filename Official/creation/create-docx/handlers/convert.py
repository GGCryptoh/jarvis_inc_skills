#!/usr/bin/env python3
"""
create-docx handler — converts Markdown to .docx via python-docx.

Reads JSON from stdin: { "markdown": "...", "filename": "doc", "title": "" }
Outputs JSON to stdout: { "result": "...", "documentUrl": "data:application/vnd...;base64,..." }
"""

import sys
import json
import re
import base64
from io import BytesIO

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print(json.dumps({"error": "python-docx not installed. Run: pip install python-docx"}))
    sys.exit(1)


def parse_inline(paragraph, text):
    """Parse inline Markdown formatting into docx runs."""
    # Pattern: **bold**, *italic*, `code`, ***bold italic***
    pattern = re.compile(r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    last_end = 0

    for match in pattern.finditer(text):
        # Add text before this match
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])

        if match.group(2):  # ***bold italic***
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.italic = True
        elif match.group(3):  # **bold**
            run = paragraph.add_run(match.group(3))
            run.bold = True
        elif match.group(4):  # *italic*
            run = paragraph.add_run(match.group(4))
            run.italic = True
        elif match.group(5):  # `code`
            run = paragraph.add_run(match.group(5))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        last_end = match.end()

    # Add remaining text
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def markdown_to_docx(md_text, title=""):
    """Convert markdown text to a python-docx Document."""
    doc = Document()

    if title:
        doc.add_heading(title, level=0)

    lines = md_text.split('\n')
    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]

        # Code block toggle
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block — add as a single paragraph with monospace
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(8)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        stripped = line.strip()

        # Empty line
        if not stripped:
            i += 1
            continue

        # Headings
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            doc.add_heading(heading_match.group(2), level=min(level, 9))
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^[-*_]{3,}$', stripped):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('_' * 40)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            i += 1
            continue

        # Unordered list
        ul_match = re.match(r'^[-*+]\s+(.+)$', stripped)
        if ul_match:
            p = doc.add_paragraph(style='List Bullet')
            parse_inline(p, ul_match.group(1))
            i += 1
            continue

        # Ordered list
        ol_match = re.match(r'^\d+[.)]\s+(.+)$', stripped)
        if ol_match:
            p = doc.add_paragraph(style='List Number')
            parse_inline(p, ol_match.group(1))
            i += 1
            continue

        # Blockquote
        bq_match = re.match(r'^>\s*(.*)$', stripped)
        if bq_match:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(36)
            run = p.add_run(bq_match.group(1))
            run.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        parse_inline(p, stripped)
        i += 1

    return doc


def main():
    try:
        raw = sys.stdin.read()
        data = json.loads(raw)
    except json.JSONDecodeError as e:
        print(json.dumps({"error": f"Invalid JSON input: {e}"}))
        sys.exit(1)

    md_text = data.get("markdown", "")
    filename = data.get("filename", "document")
    title = data.get("title", "")

    if not md_text.strip():
        print(json.dumps({"error": "No markdown content provided"}))
        sys.exit(1)

    doc = markdown_to_docx(md_text, title)

    # Save to BytesIO and encode as base64 data URI
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    b64 = base64.b64encode(buf.read()).decode('utf-8')

    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    data_uri = f"data:{mime};base64,{b64}"

    # Count document stats
    para_count = len(doc.paragraphs)
    word_count = sum(len(p.text.split()) for p in doc.paragraphs)

    safe_filename = re.sub(r'[^a-zA-Z0-9_-]', '_', filename)

    result = {
        "result": f"Generated **{safe_filename}.docx** ({para_count} paragraphs, ~{word_count} words). Download link attached.",
        "documentUrl": data_uri,
        "filename": f"{safe_filename}.docx"
    }

    print(json.dumps(result))


if __name__ == "__main__":
    main()
